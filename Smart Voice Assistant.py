# List of Library used in this project:

# Pyttsx3: pyttsx3 is a text-to-speech conversion library in Python. Unlike alternative libraries, it works offline,
# and is compatible with both Python 2 and 3.

# Datetime: The datetime module supplies classes for manipulating dates and times.

# Speech_recognition: Library for performing speech recognition, with support for several engines and APIs, online and
# offline.

# Smtplib: The smtplib module defines an SMTP client session object that can be used to send mail to any Internet
# machine with an SMTP

# Webbrowser: The webbrowser module provides a high-level interface to allow displaying Web-based documents to users.

# Psutil: Psutil (process and system utilities) is a cross-platform library for retrieving information on running
# processes and system utilization (CPU, memory, disks, network, sensors) in Python. It is useful mainly for system
# monitoring, profiling and limiting process resources and management of running processes.

# Pyjokes: One line jokes for programmers (jokes as a service)

# Os: This module provides a portable way of using operating system dependent functionality

# Pyautogui: PyAutoGUI is a cross-platform GUI automation Python module for human beings. Used to programmatically
# control the mouse & keyboard.

# Random: This module implements pseudo-random number generators for various distributions.

# Docx: Python-docx is a Python library for creating and updating Microsoft Word (.docx) files.

# Wolframalpha: Basic usage is pretty simple. Create the client with your App ID (request from Wolfram Alpha) used to
# ask queries.

# Time: This module provides various time-related functions.

# Selenium: The selenium package is used to automate web browser interaction from Python.

# Keyboard: Take full control of your keyboard with this small Python library. Hook global events, register hotkeys,
# simulate key presses and much more.

# whatsapp youtube wikipedia notes in ms word


import pyttsx3  # pip install pyttsx3 # to convert text to speech
import datetime  # to use datetime facility
import speech_recognition as sr  # pip install speech_recognition for speech recognition
import wikipedia  # pip install wikipedia
import smtplib  # pip install smtlib# To send the mail you use smtpObj to connect to the SMTP server on the local machine.
import webbrowser as wb  # to do web browser tasks
import psutil  # pip install psutil
import pyjokes  # pip install pyjokes
import os
import pyautogui  # pip install pyautogui # gui related tasks like screen shot
import random
import docx  # pip install docx
import requests
from urllib.request import urlopen
import wolframalpha  # pip install wolframalpha
import time
from selenium import webdriver  # pip install selenium
from keyboard import press  # pip install keyboard


engine = pyttsx3.init()
wolframalpha_app_id = 'Y8Q8JU-3LH8Y8XQ85'  ##'Y8Q8JU-K52G4REV2A'  #Y8Q8JU-3LH8Y8XQ85


class SystemInfo():

    def cpu(self):
        usage = str(psutil.cpu_percent())  # which return current value of cpu usage as a percentage
        speak('CPU is at' + usage)
        battery = psutil.sensors_battery()
        speak("Battery is at")
        speak(battery.percent)

    def screenshot(self):
        img = pyautogui.screenshot()
        img.save('C:\\Users\\pc\\Documents\\scrnshot\\screenshot.jpg')

    def time_(self):
        Time = datetime.datetime.now().strftime("%H:%M:%S")  # for 12 hour clock use %I for 24 hour time
        speak("The current time is")
        speak(Time)

    def date_(self):
        year = datetime.datetime.now().year
        month = datetime.datetime.now().month
        date = datetime.datetime.now().day
        speak("The current date is")
        speak(date)
        speak(month)
        speak(year)


class SystemTasks():

    def logout(self):
        os.system("shutdown -1")

    def restart(self):
        os.system("shutdown /r /t 1")

    def shutdown(self):
        os.system("shutdown /s /t 1")

    def cpu(self):
        sysinfo = SystemInfo()
        sysinfo.cpu()

    def screenshot(self):
        sysinfo = SystemInfo()
        sysinfo.screenshot()

    def time_(self):
        sysinfo = SystemInfo()
        sysinfo.time_()

    def date_(self):
        sysinfo = SystemInfo()
        sysinfo.date_()


class SearchingTasks():
    def __init__(self, obj1):
        self.obj1 = obj1

    def search(self):
        self.obj1.search()


class InternetTasks():
    def __init__(self, obj1):
        self.obj1 = obj1


def sendEmail(to, content):
    server = smtplib.SMTP('smtp.gmail.com', 587) #mailing service to use
    server.ehlo()  #help in identifing ourselves to smtp server
    server.starttls()   # which will helps us in putting connection to the smtp server into the TLS model
    # for this function, you must enable low security in your gmail which you are going to use as sender

    server.login('vivosam987@gmail.com', '9871960713')
    server.sendmail('vivosam987@gmail.com', to, content)
    server.close()


class Info():

    def __init__(self):
        self.driver = webdriver.Chrome(executable_path='C:\\Users\pc\\Drivers\\chromedriver.exe')

    def search(self, query):
        self.query = query
        self.driver.get(url="https://www.wikipedia.org/")
        search = self.driver.find_element_by_xpath('//*[@id="searchInput"]')
        search.click()
        search.send_keys(query)

        enter = self.driver.find_element_by_xpath('//*[@id="search-form"]/fieldset/button/i')
        enter.click()

        # the definition the bot can read
        info = self.driver.find_element_by_xpath('//*[@id="mw-content-text"]/div[1]/p[2]')
        readable_text = info.text
        speak(readable_text)


class Gsearch():

    def __init__(self):
        self.driver = webdriver.Chrome(executable_path='C:\\Users\\pc\\Drivers\\chromedriver.exe')

    def search(self, name):
        self.driver.get(url="https://www.google.com")
        search = self.driver.find_element_by_xpath('//*[@id="tsf"]/div[2]/div[1]/div[1]/div/div[2]/input')
        search.click()
        search.send_keys(name)

        submit = self.driver.find_element_by_name('btnK')
        submit.click()


class YouPlay():
    def __init__(self):
        self.driver = webdriver.Chrome(executable_path='C:\\Users\pc\\Drivers\\chromedriver.exe')

    def play(self, name):
        self.name = name
        self.driver.get(url="https://www.youtube.com/results?search_query=" + name)
        speak("Do you want to play first video")
        print("Do you want to play first video")
        response = TakeCommand().lower()
        if 'yes' in response or 'sure' in response:
            video = self.driver.find_element_by_xpath('//*[@id="video-title"]/yt-formatted-string')
            video.click()


class Whatsapp():

    def __init__(self):
        self.driver = webdriver.Chrome(executable_path='C:\\Users\pc\\Drivers\\chromedriver.exe')

    def send(self, reciever, message):
        self.driver.get(url="https://web.whatsapp.com/")
        self.reciever = reciever
        time.sleep(12)
        search = self.driver.find_element_by_xpath('//*[@id="side"]/div[1]/div/label/div/div[2]')
        search.click()
        search.send_keys(reciever)
        press('enter')
        text = self.driver.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[2]/div/div[2]')
        text.click()
        text.send_keys(message)
        button1 = self.driver.find_element_by_xpath('//*[@id="main"]/footer/div[1]/div[3]')
        button1.click()
        time.sleep(7)
        self.driver.close()


class StackOverflow():

    def __init__(self):
        self.driver = webdriver.Chrome(executable_path='C:\\Users\pc\\Drivers\\chromedriver.exe')

    def Search(self, query):
        self.driver.get(url="https://stackoverflow.com/")
        time.sleep(4)
        search = self.driver.find_element_by_xpath('//*[@id="search"]/div/input')
        search.click()
        search.send_keys(query)
        search.click()
        press("enter")


def speak(audio):
    engine.say(audio)
    engine.runAndWait()  # wait until function is completed


def TakeCommand():
    r = sr.Recognizer()

    with sr.Microphone() as source:
        print("Listening....")
        r.pause_threshold = 1  # wait 1 sec until user command
        audio = r.listen(source)
    try:
        print("Recognizing....")
        query = r.recognize_google(audio, language='en-US')  # audio was recognized by the google
        print(query)
    except Exception as e:
        print(e)
        print("Say again please")
        return "None"
    return query


def wishme():
    speak("Welcome back")

    # grettings
    hour = datetime.datetime.now().hour

    if hour>=6 and hour<12:
        speak("Good Morning")

    elif hour>=12 and hour<18:
        speak("Good Afternoon sir")

    elif hour>=18 and hour<24:
        speak("Good Evening sir")

    else:
        speak("Good Night sir")

    speak("Hello I am Smart Voice Assitant. Please tell me how can I help you today")


def joke():
    speak(pyjokes.get_joke())



if __name__ == "__main__":

    while True:
        query = TakeCommand().lower()  # commands convert to lowercase then stored in variable query

        if 'time' in query:
            sys1 = SystemTasks()
            sys1.time_()

        elif 'date' in query:
            sys1 = SystemTasks()
            sys1.date_()

        elif 'wishme' in query:
            wishme()

        elif 'wikipedia' in query:
            print("what do you want to search on wikipedia")
            speak("what do you want to search on wikipedia")

            try:
                query = TakeCommand().lower()
                wiki = Info()
                any1 = SearchingTasks(wiki)
                any1.obj1.search(query)
            except sr.UnknownValueError:
                print("")
            except sr.RequestError as e:
                print("")

        elif 'send email' in query:
            try:
                print("What should I say")
                speak("What should I say")
                content = TakeCommand()
                # provide reciever email address
                speak("Who is the reciever?")
                reciever = input("Enter Reciever Email : ") #abc@gmail.com
                to = reciever
                sendEmail(to, content)
                speak(content)
                print("Email has been sent")
                speak("Email has been sent")
            except Exception as e:
                print(e)
                print("Unable to send email")
                speak("Unable to send email")


        elif 'whatsapp' in query:
            print("here we go to whatsapp")
            speak("here we go to whatsapp")

            try:
                print("Speak the name of reciever")
                speak("Speak the name of reciever")
                reciever = TakeCommand().lower()#input('Enter reciever')#TakeCommand().lower()##
                print("Speak Message")
                speak("Speak Message")
                message = TakeCommand()
                Whatsapp1 = Whatsapp()
                Whatsapp1.send(reciever, message)
            except Exception as e:
                print(e)
                print("Unable to send message")
                speak("Unable to send message")

        elif 'Google meet' in query or 'Google meeting' in query:
            print("Here we go to meet")
            speak("Here we go to meet")
            speak("Enter link of meet")
            LinkOfMeet = input("Enter Link")
            wb.open(LinkOfMeet)
            time.sleep(10)
            press('ctrl+d')
            press('ctrl+e')

        elif 'meet' in query:
            print("Here we go to meet")
            speak("Here we go to meet")
            speak("Enter Link of meet")
            LinkOfMeet = input("Enter Link")
            wb.open(LinkOfMeet)

        elif 'youtube' in query:
            print("Here we go to YouTube")
            speak("Here we go to YouTube")

            try:
                print("What should i search on youtube")
                speak("What should i search on youtube")
                query = TakeCommand().lower()
                PlayOnYou = YouPlay()
                any1 = InternetTasks(PlayOnYou)
                any1.obj1.play(query)
            except sr.UnknownValueError:
                print("")
            except sr.RequestError as e:
                print("")

        elif 'stackoverflow' in query or 'stack over flow' in query or 'stack overflow' in query or 'stackover flow' in query:
            print("Here we go to StackOverFlow")
            speak("Here we go to StackOverFlow")

            try:
                print("What should i search on StackOverflow")
                speak("What should i search on StackOverflow")
                query = TakeCommand().lower()
                StackOver = StackOverflow()
                StackOver.Search(query)
            except Exception as e:
                print(e)
                print("Unable to search")
                speak("Unable to search")
        elif 'google' in query:
            print("What should I search on google")
            speak("What should I search on google")

            try:
                query = TakeCommand().lower()
                searchOnGoogle = Gsearch()
                any1 = SearchingTasks(searchOnGoogle)
                any1.obj1.search(query)
            except sr.UnknownValueError:
                print("")
            except sr.RequestError as e:
                print("")

        elif 'joke' in query or 'jokes' in query:
            joke()

        elif 'go offline' in query or 'exit' in query or 'quit' in query or 'close' in query:
            speak("Exiting the Application")
            print("Exiting the Application")
            quit()

        elif 'cleaner' in query:
            print("opening CCleaner")
            speak("opening CCleaner")
            Path = r'C:\\Program Files\\CCleaner\\CCleaner.exe'
            os.startfile(Path)

        elif 'vs code' in query or 'vscode' in query or 'visiual studio code' in query:
            print("opening VS Code")
            speak("opening VS Code")
            Path = r'C:\Users\pc\AppData\Local\Programs\Microsoft VS Code\Code.exe'
            os.startfile(Path)

        elif 'pycharm' in query or 'intellij' in query or 'python IDE' in query or 'python' in query:
            print("opening VS python IDE")
            speak("opening VS python IDE")
            Path = r'C:\\Program Files\\JetBrains\\IntelliJ IDEA Community Edition 2020.2.4\\bin\\idea64.exe'
            os.startfile(Path)

        elif 'filmora' in query or 'wondershare Filmora' in query or 'video editing software' in query or 'video edit' in query:
            print("opening Wondershare Filmora")
            speak("opening Wondershare Filmora")
            Path = r'C:\Program Files\Wondershare\Filmora\Filmora.exe'
            os.startfile(Path)

        elif 'zoom' in query:
            print("opening zoom meeting")
            speak("opening zoom meeting")
            Path = r'C:\Users\pc\AppData\Roaming\Zoom\bin\Zoom.exe'
            os.startfile(Path)

        elif 'notepad++' in query or 'notepad plus plus' in query or 'notepadplusplus' in query:
            print("opening Notepad ++")
            speak("opening Notepad ++")
            Path = r'C:\Program Files (x86)\Notepad++\notepad++.exe'
            os.startfile(Path)

        elif 'write a note' in query:
            print("what should i write, Sir?")
            speak("what should i write, Sir")
            notes = TakeCommand()
            file = open('notes.txt', 'w')
            print("Sir should i include Date and Time?")
            speak("Sir should i include Date and Time?")
            ans = TakeCommand()

            if 'yes' in ans or 'sure' in ans:
                strTime = datetime.datetime.now().strftime("%H:%M:%S")
                file.write(strTime)
                file.write(':-')
                file.write(notes)
                speak('Done Taking Notes, Sir')
            else:
                file.write(notes)

        elif 'show notes' in query:
            print("showing notes")
            speak("Showing Notes")
            file = open('notes.txt', 'r')
            print(file.read())
            speak(file.read())

        elif 'write in word' in query or 'write note in ms word' in query or 'write in ms word' in query:
            print("What is the name of file")
            speak("What is the name of file")
            fileName = TakeCommand()
            print("what is the Heading?")
            speak("what is the Heading?")
            heading = TakeCommand()
            #speak("What is the size of heading 0, 1, 2....")
            sizeOfHeading = 1 #int(TakeCommand())
            print("What should i write to notes")
            speak("What should i write to notes")
            notes = TakeCommand()
            doc = docx.Document()
            doc.add_heading(heading,sizeOfHeading)
            doc.add_paragraph(notes)
            doc.save(fileName + '.docx')

        elif 'play music' in query:
            songs_dir = 'E:\\Songs'
            music = os.listdir(songs_dir)
            print("what should i play")
            speak("what should i play")
            print("Select a number....")  # number 1
            speak("Select a number 1, 2, 3 .......")
            ans = TakeCommand().lower()

            while 'number' not in ans and ans != 'random' and ans != 'you choose' and ans != 'choose yourself':
                speak("I could not understand you. Please try again")
                ans = TakeCommand().lower()
            if 'number' in ans:  # "number 1" -> "1" -> intiger
                no = int(ans.replace('number', ''))
            elif 'random' or 'you choose' or 'choose yourself' in ans:
                no = random.randint(1, len(music))
            os.startfile(os.path.join(songs_dir, music[no]))

        elif 'remember that' in query:
            print("What should I remember?")
            speak("What should I remember?")
            memory = TakeCommand()
            print("You asked to remember that " + memory)
            speak("You asked to remember that " + memory )
            remember = open('memory.txt', 'w')
            remember.write(memory)
            remember.close()

        elif 'do you remember anything' in query:
            remember = open('memory.txt', 'r')
            print("You asked me to remember that" + remember.read())
            speak("You asked me to remember that" + remember.read())

        elif 'where is' in query:
            query = query.replace("where is", "")
            location = query
            print("User asked to locate " + location)
            speak("User asked to locate " + location)
            wb.open_new_tab("https://www.google.com/maps/place/" + location)

        elif 'calculate' in query:  # login to this site please  https://products.wolframalpha.com/api/
            client = wolframalpha.Client(wolframalpha_app_id)
            indx = query.lower().split().index('calculate')  # [calculate, sin, 45] "0"
            query = query.split()[indx + 1:]
            res = client.query(''.join(query))
            answer = next(res.results)
            #print(answer)

            try:
                print('The answer is : ' + str(answer['subpod']['img']['@title']))
                speak('The Answer is : ' + str(answer['subpod']['img']['@title']))
            except StopIteration:
                print("No Results")

        elif 'what is' in query or 'who is' in query:
            # use the same API key that we generate earliar i.e wolframalpha
            client = wolframalpha.Client(wolframalpha_app_id)
            res = client.query(query)
            answer = next(res.results)
            #print(answer)

            try:
                print(str(answer['subpod']['img']['@title']))
                speak(str(answer['subpod']['img']['@title']))
            except StopIteration:
                print("No Results")


        elif 'stop listening' in query:
            print('For How many seconds you want me to stop listening to your commands?')
            speak('For How many seconds you want me to stop listening to your commands?')
            ans = int(TakeCommand())
            time.sleep(ans)
            print(ans)

        elif 'log out' in query:
            sys1 = SystemTasks()
            sys1.logout()

        elif 'restart' in query:
            sys1 = SystemTasks()
            sys1.restart()

        elif 'shutdown' in query:
            sys1 = SystemTasks()
            sys1.shutdown()

        elif 'screenshot' in query:
            sys1 = SystemTasks()
            sys1.screenshot()
            
        elif 'cpu' in query:
            sys1.cpu()
